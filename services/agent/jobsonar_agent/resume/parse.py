from __future__ import annotations

import re
from pathlib import Path

from jobsonar_agent.resume.lexicon import SKILLS

_phrase_re = re.compile(r"[^a-z0-9+#/]+")


def normalize(text: str) -> str:
    lowered = text.lower().replace("ci / cd", "ci/cd").replace("ci-cd", "ci/cd")
    parts = _phrase_re.sub(" ", lowered)
    return " ".join(parts.split())


def extract_skills(text: str, lexicon: tuple[str, ...] = SKILLS) -> list[str]:
    hay = " " + normalize(text) + " "
    found: list[str] = []
    seen: set[str] = set()
    for skill in lexicon:
        needle = " " + normalize(skill) + " "
        if needle == "  " or needle not in hay:
            continue
        key = skill.lower()
        if key in seen:
            continue
        seen.add(key)
        found.append(skill)
    return found


def extract_text(path: str | Path) -> str:
    p = Path(path)
    suffix = p.suffix.lower()
    if suffix == ".txt":
        return p.read_text(encoding="utf-8", errors="replace")
    if suffix == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(str(p))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    if suffix == ".docx":
        from docx import Document

        doc = Document(str(p))
        parts = [para.text for para in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return "\n".join(parts)
    raise ValueError("unsupported resume type")


def parse_resume(path: str | Path) -> dict:
    text = extract_text(path)
    skills = extract_skills(text)
    return {"skills": skills, "char_count": len(text)}
